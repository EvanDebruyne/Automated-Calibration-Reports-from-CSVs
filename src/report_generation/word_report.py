"""
word_report.py
--------------
Generates a professional calibration & water quality Word report (.docx)
using python-docx.

Report sections:
  1. Cover / header (site, instrument, date range)
  2. Executive Summary (from Claude)
  3. Statistical Summary table
  4. Anomalies table
  5. Calibration Status
  6. Corrective Actions
  7. Compliance Notes
  8. Raw data appendix (optional)
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Logo path — relative to project root
LOGO_PATH = Path(__file__).parent.parent.parent / "templates" / "aquasummit_logo.png"

from src.claude_analysis.analyzer import AnalysisResult

logger = logging.getLogger(__name__)

# ── Unit conversion helpers ───────────────────────────────────────────────────
# Display labels shown in reports (imperial units)
_COL_LABELS = {
    "ph":               "pH",
    "turbidity_ntu":    "Turbidity (NTU)",
    "flow_rate_lps":    "Flow Rate (GPM)",
    "temperature_c":    "Temperature (°F)",
    "conductivity_us":  "Conductivity (µS/cm)",
}

def _display_col(col: str) -> str:
    return _COL_LABELS.get(col, col.replace("_", " ").title())

def _convert_units(df: pd.DataFrame) -> pd.DataFrame:
    """Return a copy of df with temperature in °F and flow in GPM.
    Unknown numeric columns are kept as-is with a cleaned-up display name."""
    d = df.copy()
    if "temperature_c" in d.columns:
        d["temperature_c"] = (d["temperature_c"] * 9 / 5 + 32).round(2)
    if "flow_rate_lps" in d.columns:
        d["flow_rate_lps"] = (d["flow_rate_lps"] * 15.8503).round(2)
    return d.rename(columns=_COL_LABELS)

def _get_numeric_cols(df: pd.DataFrame) -> list[str]:
    """Return all numeric sensor columns, excluding metadata fields."""
    skip = {"site_id", "instrument_id", "timestamp"}
    return [c for c in df.select_dtypes(include="number").columns if c not in skip]

# Brand colours
COLOUR_HEADER = RGBColor(0x1A, 0x57, 0x8A)   # Deep blue
COLOUR_TABLE_HEADER = RGBColor(0x2E, 0x75, 0xB6)
COLOUR_CRITICAL = RGBColor(0xC0, 0x00, 0x00)
COLOUR_HIGH = RGBColor(0xFF, 0x00, 0x00)
COLOUR_MEDIUM = RGBColor(0xFF, 0xA5, 0x00)
COLOUR_LOW = RGBColor(0x00, 0x70, 0xC0)

SEVERITY_COLOURS = {
    "CRITICAL": "C00000",
    "HIGH": "FF0000",
    "MEDIUM": "FFA500",
    "LOW": "0070C0",
}


class WordReportGenerator:
    """
    Builds a .docx water quality report.

    Usage
    -----
    >>> gen = WordReportGenerator()
    >>> path = gen.generate(
    ...     analysis_result=result,
    ...     df=df,
    ...     site_id="SITE-01",
    ...     instrument_id="FT-201",
    ...     output_path="examples/generated_reports/report.docx",
    ... )
    """

    def generate(
        self,
        analysis_result: AnalysisResult,
        df: pd.DataFrame,
        site_id: str,
        instrument_id: str,
        output_path: str | Path,
        technician_name: str = "Field Technician",
        include_raw_data: bool = False,
    ) -> Path:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        self._set_margins(doc)
        self._apply_styles(doc)

        self._add_header(doc, site_id, instrument_id, df, analysis_result)
        self._add_executive_summary(doc, analysis_result)
        self._add_statistical_summary(doc, df)
        self._add_anomalies(doc, analysis_result)
        self._add_calibration_service_notes(doc, analysis_result)
        self._add_footer_info(doc, technician_name)

        if include_raw_data:
            self._add_raw_data_appendix(doc, df)

        doc.save(str(output_path))
        logger.info("Word report saved: %s", output_path)
        return output_path

    # ── Document sections ─────────────────────────────────────────────────────

    def _add_header(self, doc: Document, site_id: str, instrument_id: str, df: pd.DataFrame, result: AnalysisResult):
        # Logo + title side by side using a borderless 2-col table
        logo_table = doc.add_table(rows=1, cols=2)
        logo_table.style = "Table Grid"
        # Remove all borders on the logo table
        for row in logo_table.rows:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                tcBorders = OxmlElement("w:tcBorders")
                for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                    border = OxmlElement(f"w:{side}")
                    border.set(qn("w:val"), "nil")
                    tcBorders.append(border)
                tc_pr.append(tcBorders)

        # Left cell: logo image
        logo_cell = logo_table.rows[0].cells[0]
        logo_para = logo_cell.paragraphs[0]
        logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if LOGO_PATH.exists():
            run = logo_para.add_run()
            run.add_picture(str(LOGO_PATH), width=Inches(2.2))

        # Right cell: report title
        title_cell = logo_table.rows[0].cells[1]
        title_para = title_cell.paragraphs[0]
        title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        title_run = title_para.add_run("Water Quality Calibration Report")
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = COLOUR_HEADER

        doc.add_paragraph()  # spacer

        # Date range
        if "timestamp" in df.columns:
            date_from = pd.to_datetime(df["timestamp"]).min().strftime("%d %B %Y")
            date_to = pd.to_datetime(df["timestamp"]).max().strftime("%d %B %Y")
        else:
            date_from = date_to = "N/A"

        # Info table: site | instrument | period | status
        # Column widths in DXA: total 9360. Period gets more room for the date range.
        info_col_widths = [1800, 1800, 3600, 2160]

        STATUS_COLOURS_HEX = {
            "OK": "00B050", "DUE": "FFA500", "OVERDUE": "FF0000",
            "SUSPECTED_DRIFT": "C00000", "UNKNOWN": "7F7F7F",
        }
        status = result.calibration_status or "UNKNOWN"
        status_hex = STATUS_COLOURS_HEX.get(status, "7F7F7F")

        info_table = doc.add_table(rows=2, cols=4)
        info_table.style = "Table Grid"
        info_table.autofit = False
        labels = ["Site", "Instrument", "Report Period", "Calibration Status"]
        values = [site_id, instrument_id, f"{date_from} — {date_to}", status]

        for col_idx, (label, value) in enumerate(zip(labels, values)):
            label_cell = info_table.rows[0].cells[col_idx]
            label_cell.text = label
            self._shade_cell(label_cell, "2E75B6")
            self._set_cell_font(label_cell, bold=True, colour="FFFFFF")
            self._set_cell_width(label_cell, info_col_widths[col_idx])
            label_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            value_cell = info_table.rows[1].cells[col_idx]
            value_cell.text = value
            self._shade_cell(value_cell, "EBF3FB")
            self._set_cell_width(value_cell, info_col_widths[col_idx])
            value_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if col_idx == 3:
                self._set_cell_font(value_cell, bold=True, colour=status_hex)

        generated = doc.add_paragraph()
        generated.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = generated.add_run(f"Generated: {datetime.now().strftime('%d %B %Y %H:%M')}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

        doc.add_paragraph()  # spacer

    def _add_executive_summary(self, doc: Document, result: AnalysisResult):
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(result.executive_summary or "No summary available.")

    def _add_statistical_summary(self, doc: Document, df: pd.DataFrame):
        doc.add_heading("Statistical Summary", level=1)

        numeric_cols = _get_numeric_cols(df)
        if not numeric_cols:
            doc.add_paragraph("No numeric sensor data available.")
            return

        display_df = _convert_units(df)
        display_cols = [_display_col(c) for c in numeric_cols]

        stats = display_df[[_display_col(c) for c in numeric_cols]].describe().loc[["mean", "min", "max", "std"]].round(2)
        n_cols = len(numeric_cols) + 1

        table = doc.add_table(rows=1 + len(stats), cols=n_cols)
        table.style = "Table Grid"

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Statistic"
        self._shade_cell(hdr_cells[0], "2E75B6")
        self._set_cell_font(hdr_cells[0], bold=True, colour="FFFFFF")

        for i, col in enumerate(display_cols, start=1):
            hdr_cells[i].text = col
            self._shade_cell(hdr_cells[i], "2E75B6")
            self._set_cell_font(hdr_cells[i], bold=True, colour="FFFFFF")

        # Data rows
        stat_labels = {"mean": "Mean", "min": "Minimum", "max": "Maximum", "std": "Std Dev"}
        for r_idx, (stat_name, row) in enumerate(stats.iterrows(), start=1):
            cells = table.rows[r_idx].cells
            cells[0].text = stat_labels.get(stat_name, stat_name)
            self._set_cell_font(cells[0], bold=True)
            shade = "D9E8F5" if r_idx % 2 == 0 else "FFFFFF"
            self._shade_cell(cells[0], shade)
            for c_idx, col in enumerate(display_cols, start=1):
                cells[c_idx].text = str(row[col])
                self._shade_cell(cells[c_idx], shade)

        doc.add_paragraph()

    def _add_anomalies(self, doc: Document, result: AnalysisResult):
        doc.add_heading("Detected Anomalies", level=1)

        if not result.anomalies:
            doc.add_paragraph("No anomalies detected in the review period.")
            return

        # Column widths in DXA (1440 = 1 inch). Total = 9360 (6.5" content width)
        # Timestamp=2340, Parameter=2340, Value=1800, Severity=2880
        col_widths = [2340, 2340, 1800, 2880]

        table = doc.add_table(rows=1 + len(result.anomalies), cols=4)
        table.style = "Table Grid"
        table.autofit = False

        # Set overall table width
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OxmlElement
        tbl_pr = table._tbl.tblPr
        tbl_w = _OxmlElement("w:tblW")
        tbl_w.set(_qn("w:w"), "9360")
        tbl_w.set(_qn("w:type"), "dxa")
        tbl_pr.append(tbl_w)

        headers = ["Timestamp", "Parameter", "Value", "Severity"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            self._shade_cell(cell, "2E75B6")
            self._set_cell_font(cell, bold=True, colour="FFFFFF")
            self._set_cell_width(cell, col_widths[i])

        for r_idx, anomaly in enumerate(result.anomalies, start=1):
            cells = table.rows[r_idx].cells
            cells[0].text = str(anomaly.get("timestamp", ""))
            cells[1].text = anomaly.get("parameter", "").replace("_", " ").title()
            cells[2].text = str(anomaly.get("value", ""))
            severity = anomaly.get("severity", "LOW")
            cells[3].text = severity
            sev_colour = SEVERITY_COLOURS.get(severity, "000000")
            self._set_cell_font(cells[3], bold=True, colour=sev_colour)
            shade = "FFF2CC" if severity in ("HIGH", "CRITICAL") else "FFFFFF"
            for i, cell in enumerate(cells):
                self._shade_cell(cell, shade)
                self._set_cell_width(cell, col_widths[i])

        doc.add_paragraph()

    def _add_calibration_service_notes(self, doc: Document, result: AnalysisResult):
        doc.add_heading("Calibration & Service Notes", level=1)
        if not result.corrective_actions:
            doc.add_paragraph("No service actions recorded.")
            return
        for action in result.corrective_actions:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(action)

    def _add_footer_info(self, doc: Document, technician_name: str):
        doc.add_paragraph()
        doc.add_heading("Report Sign-Off", level=1)
        for label in ["Prepared by:", "Reviewed by:", "Date:"]:
            p = doc.add_paragraph()
            p.add_run(f"{label}  ").bold = True
            underline_space = "_" * 40
            p.add_run(underline_space if label != "Prepared by:" else f"  {technician_name}")

    def _add_raw_data_appendix(self, doc: Document, df: pd.DataFrame):
        doc.add_page_break()
        doc.add_heading("Appendix — Raw Sensor Data", level=1)
        cols = df.columns.tolist()
        table = doc.add_table(rows=1 + min(len(df), 500), cols=len(cols))
        table.style = "Table Grid"

        for i, col in enumerate(cols):
            cell = table.rows[0].cells[i]
            cell.text = col
            self._shade_cell(cell, "2E75B6")
            self._set_cell_font(cell, bold=True, colour="FFFFFF")

        for r_idx, (_, row) in enumerate(df.head(500).iterrows(), start=1):
            for c_idx, col in enumerate(cols):
                table.rows[r_idx].cells[c_idx].text = str(row[col])

    # ── Styling helpers ───────────────────────────────────────────────────────

    @staticmethod
    def _set_margins(doc: Document, margin_inches: float = 1.0):
        for section in doc.sections:
            section.top_margin = Inches(margin_inches)
            section.bottom_margin = Inches(margin_inches)
            section.left_margin = Inches(margin_inches)
            section.right_margin = Inches(margin_inches)

    @staticmethod
    def _apply_styles(doc: Document):
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

    @staticmethod
    def _shade_cell(cell, fill_hex: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)
        tc_pr.append(shd)

    @staticmethod
    def _set_cell_width(cell, width_dxa: int):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = OxmlElement("w:tcW")
        tc_w.set(qn("w:w"), str(width_dxa))
        tc_w.set(qn("w:type"), "dxa")
        tc_pr.append(tc_w)

    @staticmethod
    def _set_cell_font(cell, bold: bool = False, colour: Optional[str] = None):
        for para in cell.paragraphs:
            for run in para.runs:
                if bold:
                    run.bold = True
                if colour:
                    run.font.color.rgb = RGBColor(
                        int(colour[:2], 16), int(colour[2:4], 16), int(colour[4:], 16)
                    )

